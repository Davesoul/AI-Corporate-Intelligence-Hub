"""RAG Manager - Document ingestion and retrieval for PDFs and other files."""

import os
import pickle
from pathlib import Path
from typing import List, Optional
from pypdf import PdfReader
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlmodel import Session, select
from db import engine
from models import UploadedFile

# Directory to store uploaded documents and vector store
UPLOAD_DIR = Path("uploads")
VECTOR_STORE_PATH = Path("vector_store.pkl")
UPLOAD_DIR.mkdir(exist_ok=True)

# Initialize embeddings model (using a small, fast model)
embeddings = None
documents_store: List[Document] = []
document_embeddings: List[List[float]] = []


def get_embeddings():
    """Lazy load embeddings model."""
    global embeddings
    if embeddings is None:
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"}
        )
    return embeddings


def load_vector_store():
    """Load existing vector store from disk."""
    global documents_store, document_embeddings
    if VECTOR_STORE_PATH.exists():
        with open(VECTOR_STORE_PATH, "rb") as f:
            data = pickle.load(f)
            documents_store = data.get("documents", [])
            document_embeddings = data.get("embeddings", [])


def save_vector_store():
    """Save vector store to disk."""
    with open(VECTOR_STORE_PATH, "wb") as f:
        pickle.dump({
            "documents": documents_store,
            "embeddings": document_embeddings
        }, f)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file types."""
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".txt", ".md", ".csv", ".json", ".py", ".js", ".html", ".css"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext in [".doc", ".docx"]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return "DOCX support requires python-docx library"
    else:
        return f"Unsupported file type: {ext}"


def ingest_document(file_path: str, metadata: Optional[dict] = None, uploaded_by: Optional[int] = None) -> dict:
    """Ingest a document into the RAG system and save to database."""
    global documents_store, document_embeddings
    
    try:
        path = Path(file_path)
        
        # Extract text
        text = extract_text_from_file(file_path)
        if not text.strip():
            return {"error": "No text content found in document"}
        
        # Split into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        
        # Create documents with metadata
        file_name = path.name
        base_metadata = {"source": file_name, "file_path": file_path}
        if metadata:
            base_metadata.update(metadata)
        
        docs = [
            Document(page_content=chunk, metadata={**base_metadata, "chunk": i})
            for i, chunk in enumerate(chunks)
        ]
        
        # Generate embeddings
        emb_model = get_embeddings()
        new_embeddings = emb_model.embed_documents([doc.page_content for doc in docs])
        
        # Add to store
        documents_store.extend(docs)
        document_embeddings.extend(new_embeddings)
        
        # Save to disk
        save_vector_store()
        
        # Save to database
        file_size = path.stat().st_size if path.exists() else 0
        uploaded_file = UploadedFile(
            filename=file_name,
            file_path=str(path.absolute()),
            file_size=file_size,
            file_type=path.suffix.lower(),
            chunks_count=len(chunks),
            uploaded_by=uploaded_by
        )
        with Session(engine) as session:
            session.add(uploaded_file)
            session.commit()
        
        return {
            "status": "success",
            "file": file_name,
            "chunks_created": len(chunks),
            "total_documents": len(documents_store)
        }
    except Exception as e:
        return {"error": str(e)}


def similarity_search(query: str, k: int = 5) -> List[Document]:
    """Search for similar documents using cosine similarity."""
    # Always reload vector store to get latest documents
    load_vector_store()
    
    if not documents_store or not document_embeddings:
        return []
    
    try:
        import numpy as np
        
        # Get query embedding
        emb_model = get_embeddings()
        query_embedding = emb_model.embed_query(query)
        
        # Calculate cosine similarities
        query_np = np.array(query_embedding)
        doc_embeddings_np = np.array(document_embeddings)
        
        # Normalize vectors
        query_norm = query_np / np.linalg.norm(query_np)
        doc_norms = doc_embeddings_np / np.linalg.norm(doc_embeddings_np, axis=1, keepdims=True)
        
        # Calculate similarities
        similarities = np.dot(doc_norms, query_norm)
        
        # Get top k indices
        top_indices = np.argsort(similarities)[-k:][::-1]
        
        # Lower threshold to 0.1 for better recall
        return [documents_store[i] for i in top_indices if similarities[i] > 0.1]
    except Exception as e:
        print(f"Search error: {e}")
        return []


def query_documents(query: str, k: int = 5) -> str:
    """Query the document store and return formatted results."""
    print(f"[RAG] Searching for: {query}")
    
    # Reload to ensure we have latest
    load_vector_store()
    print(f"[RAG] Vector store has {len(documents_store)} chunks")
    
    results = similarity_search(query, k)
    print(f"[RAG] Found {len(results)} relevant chunks")
    
    if not results:
        # Check if any documents exist at all
        with Session(engine) as session:
            statement = select(UploadedFile).where(UploadedFile.is_active == True)
            files = session.exec(statement).all()
            if files:
                return f"No relevant content found for your query. Available documents: {', '.join([f.filename for f in files])}. Try a different search term."
        return "No documents have been uploaded yet. Please upload documents first using the upload feature."
    
    response = ""
    for i, doc in enumerate(results, 1):
        source = doc.metadata.get("source", "Unknown")
        content = doc.page_content
        response += f"[Source: {source}]\n{content}\n\n---\n\n"
    
    return response


def list_ingested_documents() -> List[dict]:
    """List all ingested documents from the database."""
    # Reload vector store to get current chunk count
    load_vector_store()
    
    with Session(engine) as session:
        statement = select(UploadedFile).where(UploadedFile.is_active == True)
        files = session.exec(statement).all()
        return [
            {
                "id": f.id,
                "filename": f.filename,
                "file_path": f.file_path,
                "file_size": f.file_size,
                "file_type": f.file_type,
                "chunks": f.chunks_count,
                "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None,
                "vector_store_total": len(documents_store)
            }
            for f in files
        ]


def clear_documents() -> dict:
    """Clear all ingested documents from vector store and database."""
    global documents_store, document_embeddings
    documents_store = []
    document_embeddings = []
    
    if VECTOR_STORE_PATH.exists():
        os.remove(VECTOR_STORE_PATH)
    
    # Mark all files as inactive in database
    with Session(engine) as session:
        statement = select(UploadedFile).where(UploadedFile.is_active == True)
        files = session.exec(statement).all()
        for f in files:
            f.is_active = False
        session.commit()
    
    return {"status": "cleared", "message": "All documents removed from RAG system"}


# Load existing vector store on module import
load_vector_store()
